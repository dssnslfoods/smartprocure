from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ─── Color constants ─────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
AMBER  = RGBColor(0xD9, 0x77, 0x06)
RED    = RGBColor(0xDC, 0x26, 0x26)
GRAY   = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x11, 0x18, 0x27)
LIGHT  = RGBColor(0xF1, 0xF5, 0xF9)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = 'TH Sarabun New'
    if color:
        run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = BLUE
    return p

def add_para(doc, text, bold=False, italic=False, size=12, color=None, indent=False, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_step(doc, n, text, detail=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{n}.  ')
    r1.font.name = 'TH Sarabun New'
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.font.name = 'TH Sarabun New'
    r2.font.size = Pt(11)
    if detail:
        r3 = p.add_run(f'  — {detail}')
        r3.font.name = 'TH Sarabun New'
        r3.font.size = Pt(10)
        r3.font.italic = True
        r3.font.color.rgb = GRAY
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'💡  {text}')
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = AMBER
    return p

def add_warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'⚠️  {text}')
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RED
    return p

def add_table(doc, headers, rows, col_widths=None, alt_rows=True):
    n_cols = len(headers)
    t = doc.add_table(rows=1, cols=n_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'TH Sarabun New'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        set_cell_bg(cell, '1E3A5F')

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = 'F1F5F9' if (alt_rows and ri % 2 == 1) else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            set_cell_bg(cell, bg)

    if col_widths:
        set_col_widths(t, col_widths)

    # space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

# Fix import for page break
import docx.enum.text

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SmartProcure')
r.font.name = 'TH Sarabun New'
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('ระบบจัดการจัดซื้อจัดจ้างอัจฉริยะ')
r2.font.name = 'TH Sarabun New'
r2.font.size = Pt(20)
r2.font.color.rgb = BLUE

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('คู่มือ Workflow การทำงานของระบบ')
r3.font.name = 'TH Sarabun New'
r3.font.size = Pt(24)
r3.font.bold = True
r3.font.color.rgb = NAVY

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('สำหรับทีมจัดซื้อและผู้ที่เกี่ยวข้อง')
r4.font.name = 'TH Sarabun New'
r4.font.size = Pt(14)
r4.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('NSL Foods PLC')
r5.font.name = 'TH Sarabun New'
r5.font.size = Pt(14)
r5.font.bold = True
r5.font.color.rgb = NAVY

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run('เอกสารใช้ภายใน  |  เวอร์ชัน 1.0  |  เมษายน 2026')
r6.font.name = 'TH Sarabun New'
r6.font.size = Pt(11)
r6.font.color.rgb = GRAY

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'สารบัญ', level=2)
toc_items = [
    ('1.', 'ภาพรวมระบบ SmartProcure'),
    ('2.', 'วงจรที่ 1 — Master Catalog & ราคาอ้างอิง'),
    ('3.', 'วงจรที่ 2 — Supplier เสนอราคา (Price Submission)'),
    ('4.', 'วงจรที่ 3 — สร้าง RFQ & รับ Quotation'),
    ('5.', 'วงจรที่ 4 — ประเมินผลและ Award'),
    ('6.', 'โมดูลสนับสนุน (Supplier Management & Admin Settings)'),
    ('7.', 'สิทธิ์การใช้งานแต่ละ Role'),
    ('8.', 'Quick Reference — สรุป Flow แบบย่อ'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num}  ')
    r1.font.bold = True
    r1.font.color.rgb = BLUE
    r1.font.name = 'TH Sarabun New'
    r1.font.size = Pt(11)
    r2 = p.add_run(title)
    r2.font.name = 'TH Sarabun New'
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. ภาพรวมระบบ SmartProcure', level=1)
add_para(doc, 'SmartProcure คือระบบบริหารจัดการจัดซื้อจัดจ้างครบวงจรของ NSL Foods PLC ออกแบบมาเพื่อรองรับมาตรฐาน Food Safety (BRCGS) ควบคู่กับการควบคุมต้นทุนและความโปร่งใสในกระบวนการจัดซื้อ')

add_heading(doc, '4 วงจรหลักของระบบ', level=3)
add_table(doc,
    ['วงจร', 'ชื่อ', 'ผู้รับผิดชอบหลัก', 'ผลลัพธ์'],
    [
        ['วงจรที่ 1', 'Master Catalog & ราคาอ้างอิง', 'Admin / Procurement Officer', 'Catalog พร้อมใช้'],
        ['วงจรที่ 2', 'Supplier เสนอราคา', 'Supplier', 'ราคาบันทึกใน Database'],
        ['วงจรที่ 3', 'RFQ & รับ Quotation', 'Procurement + Supplier', 'Quotation ครบถ้วน'],
        ['วงจรที่ 4', 'ประเมินผลและ Award', 'Procurement / QA', 'Purchase Order'],
    ],
    col_widths=[2.5, 5, 5, 4]
)

add_heading(doc, 'ผู้ใช้งานในระบบ (Roles)', level=3)
add_table(doc,
    ['Role', 'ชื่อในระบบ', 'หน้าที่หลัก'],
    [
        ['Admin', 'admin', 'ดูแลระบบ, จัดการ User, กำหนด Settings'],
        ['จัดซื้อ', 'procurement_officer', 'สร้าง Catalog, RFQ, ประเมินผล, Award'],
        ['Supplier', 'supplier', 'เสนอราคา, Submit Quotation, แจ้ง Decline'],
        ['QA', 'qa', 'ประเมินความเสี่ยง, กำหนด Nominated Supplier, ร่วมประเมิน'],
    ],
    col_widths=[3, 4, 10]
)
add_note(doc, 'Nominated Supplier คือ Supplier ที่ QA กำหนดให้จัดซื้อต้องซื้อจากเจ้านั้นเท่านั้น ตามข้อกำหนด BRCGS Food Safety')

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — MASTER CATALOG
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. วงจรที่ 1 — Master Catalog & ราคาอ้างอิง', level=1)
add_para(doc, 'Master Catalog คือฐานข้อมูลสินค้าหลักของระบบ จัดซื้อใช้เป็นจุดเริ่มต้นในการกำหนดความต้องการและราคาอ้างอิงก่อนออก RFQ')

add_heading(doc, 'หมวดหมู่ Catalog (4 หมวด)', level=3)
add_table(doc,
    ['Catalog', 'รายการตัวอย่าง'],
    [
        ['วัตถุดิบบรรจุภัณฑ์', 'กล่อง, ฟิล์ม, ถุง, ฉลาก, กระดาษรอง'],
        ['วัตถุดิบอาหาร', 'แป้ง, น้ำมัน, น้ำตาล, เครื่องเทศ, สารปรุงแต่งรส'],
        ['เคมีภัณฑ์ & Cleaning', 'สารทำความสะอาด, น้ำยาฆ่าเชื้อ, สารซักล้าง'],
        ['อะไหล่ & MRO', 'อะไหล่เครื่องจักร, อุปกรณ์ซ่อมบำรุง, สารหล่อลื่น'],
    ],
    col_widths=[6, 11]
)

add_heading(doc, 'ขั้นตอนการตั้งค่า Catalog', level=3)
add_step(doc, 1, 'Login ด้วย Role Admin หรือ Procurement Officer')
add_step(doc, 2, 'ไปที่เมนู Pricelist → เลือก Catalog ที่ต้องการ')
add_step(doc, 3, 'ตรวจสอบรายการสินค้า — แต่ละรายการมีข้อมูลดังนี้:')
add_bullet(doc, 'ชื่อสินค้า และ Specification')
add_bullet(doc, 'ประเภท: Open (แข่งราคา) หรือ Nominated (กำหนด Supplier)')
add_bullet(doc, 'ปริมาณเป้าหมาย (Target Qty) — จำนวนที่วางแผนสั่งซื้อ')
add_bullet(doc, 'ราคาอ้างอิง (Baseline) — คำนวณอัตโนมัติ')
add_step(doc, 4, 'แก้ไข Target Qty', 'คลิกที่ตัวเลขในคอลัมน์ Target Qty แล้วพิมพ์จำนวนใหม่')

add_heading(doc, 'ประเภทรายการสินค้า', level=3)
add_table(doc,
    ['ประเภท', 'ความหมาย', 'ใครเสนอราคาได้', 'การคำนวณ Baseline'],
    [
        ['Open', 'เปิดแข่งราคาได้ทั่วไป', 'Supplier ที่ได้รับอนุมัติทุกราย', 'ค่าเฉลี่ยทุก Offer'],
        ['Nominated', 'QA กำหนด Supplier ไว้ (BRCGS)', 'เฉพาะ Supplier ที่ถูก Nominate', 'ราคาของ Nominated Supplier'],
    ],
    col_widths=[3, 5, 5, 5]
)
add_warning(doc, 'รายการ Nominated — จัดซื้อไม่สามารถเปลี่ยน Supplier ได้โดยไม่ผ่านการอนุมัติ QA เนื่องจากเป็นข้อกำหนดด้าน BRCGS Food Safety')

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PRICE SUBMISSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. วงจรที่ 2 — Supplier เสนอราคา', level=1)
add_para(doc, 'Supplier สามารถเสนอราคาได้ 2 ช่องทาง ได้แก่ Excel Checklist (Offline) และ Supplier Portal (Online)')

add_heading(doc, '3A — ช่องทาง Excel Checklist (Offline)', level=2)
add_heading(doc, 'ฝั่งจัดซื้อ — ส่ง Checklist', level=3)
add_step(doc, 1, 'ไปที่ Pricelist → เลือก Catalog')
add_step(doc, 2, 'ใช้ Checkbox เลือกรายการสินค้าที่ต้องการขอราคา')
add_step(doc, 3, 'เลือก Supplier จาก Dropdown และกรอก RFQ Number (ถ้ามี)')
add_step(doc, 4, 'กดปุ่ม "ส่ง Excel Checklist" → ไฟล์ .xlsx ถูก Download')
add_step(doc, 5, 'ส่งไฟล์ให้ Supplier ทาง Email หรือช่องทางอื่น')

add_heading(doc, 'ฝั่ง Supplier — กรอกราคา', level=3)
add_step(doc, 1, 'เปิดไฟล์ Excel ที่ได้รับ')
add_para(doc, 'ไฟล์มี 2 Sheet:', italic=True, color=GRAY, indent=True)
add_bullet(doc, 'Sheet "Info" — ข้อมูล Supplier, วันที่, RFQ Number')
add_bullet(doc, 'Sheet "Checklist" — รายการสินค้าพร้อมคอลัมน์ "ราคาต่อหน่วย" ให้กรอก')
add_step(doc, 2, 'กรอกราคาในคอลัมน์ "ราคาต่อหน่วย" (บาท)')
add_step(doc, 3, 'Save ไฟล์และส่งกลับให้จัดซื้อ')
add_note(doc, 'ไม่ควรเปลี่ยนแปลงคอลัมน์อื่น โดยเฉพาะคอลัมน์ ID ที่ซ่อนอยู่ ระบบใช้ ID นี้จับคู่ข้อมูล')

add_heading(doc, 'ฝั่งจัดซื้อ — Import ราคากลับ', level=3)
add_step(doc, 1, 'ในหน้า Catalog → กดปุ่ม "Import Quotation"')
add_step(doc, 2, 'เลือกไฟล์ Excel ที่ Supplier ส่งกลับมา')
add_step(doc, 3, 'ระบบแสดง Preview รายการก่อน Confirm')
add_step(doc, 4, 'กด "Confirm" → ราคาบันทึกเข้า Database ทันที')

add_heading(doc, '3B — ช่องทาง Supplier Portal (Online)', level=2)
add_heading(doc, 'Supplier Login เข้าระบบ', level=3)
add_table(doc,
    ['ข้อมูล Login', 'รายละเอียด'],
    [
        ['Username', 'sup-XXXXX@nslfoods.test  (แทนที่ XXXXX ด้วย Supplier Code)'],
        ['Password', '123456  (เปลี่ยนได้ในระบบ)'],
    ],
    col_widths=[5, 12]
)

add_heading(doc, 'ขั้นตอนเสนอราคาใน Portal', level=3)
add_step(doc, 1, 'Login → ไปที่เมนู Pricelist')
add_step(doc, 2, 'เลือก Catalog ที่ต้องการเสนอราคา')
add_step(doc, 3, 'กรอกราคาในคอลัมน์ "ราคาต่อหน่วย"')
add_note(doc, 'Supplier เห็นเฉพาะรายการ Open ทั้งหมด และรายการ Nominated ที่ตัวเองถูก Nominate')
add_step(doc, 4, 'กด Submit → จัดซื้อเห็นราคาทันที')

add_heading(doc, 'รอบการเสนอราคา (Update Cycle)', level=3)
add_table(doc,
    ['สถานะ', 'ความหมาย', 'สัญลักษณ์'],
    [
        ['ส่งแล้ว (Fresh)', 'ยังอยู่ในช่วง Hold Period ไม่ต้องส่งใหม่', '🟢 สีเขียว'],
        ['ใกล้ถึงรอบ (Due Soon)', 'เหลือน้อยกว่า 30 วันก่อนถึงรอบ', '🟡 สีเหลือง'],
        ['ถึงรอบแล้ว (Overdue)', 'เกินกำหนดแล้ว ควรส่งราคาใหม่', '🔴 สีแดง'],
        ['ยังไม่เคยส่ง (Never)', 'ยังไม่มีการเสนอราคาในรอบนี้', '⚫ สีเทา'],
    ],
    col_widths=[5, 8, 4]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — RFQ
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. วงจรที่ 3 — สร้าง RFQ & รับ Quotation', level=1)

add_heading(doc, '4A — การสร้าง RFQ', level=2)

add_heading(doc, 'วิธีที่ 1: สร้างจาก Price List (แนะนำ)', level=3)
add_step(doc, 1, 'ในหน้า Catalog → เลือกรายการที่ต้องการสั่งซื้อด้วย Checkbox')
add_step(doc, 2, 'กดปุ่ม "Create RFQ"')
add_step(doc, 3, 'ระบบสร้าง RFQ + รายการสินค้า (rfq_items) โดยอัตโนมัติ')
add_step(doc, 4, 'Supplier ที่เคยเสนอราคารายการนั้นถูกเพิ่มเป็น Invited Suppliers อัตโนมัติ')
add_note(doc, 'วิธีนี้ช่วยประหยัดเวลา เพราะระบบเชื่อมโยงข้อมูลราคาอ้างอิงจาก Catalog เข้า RFQ ให้อัตโนมัติ')

add_heading(doc, 'วิธีที่ 2: สร้างเอง (Manual)', level=3)
add_step(doc, 1, 'ไปที่เมนู RFQ → กดปุ่ม "New RFQ"')
add_step(doc, 2, 'กรอกข้อมูล: ชื่อ RFQ, วันสิ้นสุด, หมวดหมู่')
add_step(doc, 3, 'เพิ่มรายการสินค้าในแท็บ Items')
add_step(doc, 4, 'เชิญ Supplier ในแท็บ Invite Suppliers')

add_heading(doc, '4B — สถานะ RFQ', level=2)
add_para(doc, 'สถานะ RFQ เปลี่ยนได้จาก Dropdown (ไม่ใช่ปุ่มลำดับแบบ Step)')
add_table(doc,
    ['สถานะ', 'ความหมาย', 'เงื่อนไขในการเปลี่ยน'],
    [
        ['Draft', 'กำลังเตรียม RFQ', 'ไม่มีเงื่อนไข'],
        ['Published', 'ส่งให้ Supplier แล้ว รับ Quotation ได้', 'ต้องมี Supplier ≥ 1 รายในรายการ'],
        ['Closed', 'ปิดรับ Quotation แล้ว', 'ไม่มีเงื่อนไข'],
        ['Evaluation', 'อยู่ระหว่างประเมินราคา', 'ไม่มีเงื่อนไข'],
        ['Awarded', 'มอบหมายงานแล้ว', 'ต้องมีการเลือก Recommended Winner ก่อน'],
    ],
    col_widths=[3.5, 7, 7]
)

add_heading(doc, '4C — การเชิญ Supplier', level=2)
add_heading(doc, 'ส่วน "Invited" — Supplier ที่เชิญแล้ว', level=3)
add_bullet(doc, 'แสดงสถานะ: รอตอบ / ตอบแล้ว / ถอนตัว (พร้อมเหตุผล)')
add_bullet(doc, 'จัดซื้อสามารถ Remove Supplier ออกจาก RFQ ได้ (เฉพาะตอน Status = Draft)')

add_heading(doc, 'ส่วน "Available" — Supplier ที่ยังไม่ได้เชิญ', level=3)
add_table(doc,
    ['Eligibility Badge', 'ความหมาย'],
    [
        ['✅  Eligible', 'เชิญได้ทันที ผ่านเกณฑ์ทุกข้อ'],
        ['⚠️  Warning', 'มีข้อควรระวัง เช่น ใบรับรองใกล้หมดอายุ'],
        ['🔴  Blocked', 'ถูก Blacklist หรือ Status ไม่ผ่าน ไม่ควรเชิญ'],
    ],
    col_widths=[5, 12]
)

add_heading(doc, '4D — Supplier รับ / ปฏิเสธ RFQ', level=2)

add_heading(doc, 'กรณีที่ 1 — เสนอราคา (Submit Quotation)', level=3)
add_step(doc, 1, 'Supplier Login → ไปที่ RFQ ที่ได้รับเชิญ')
add_step(doc, 2, 'เปิดแท็บ "Quotations"')
add_step(doc, 3, 'เห็นชื่อบริษัทตัวเองแสดงอัตโนมัติ (ไม่ต้องเลือก)')
add_step(doc, 4, 'กรอกราคาแต่ละรายการ')
add_step(doc, 5, 'กด Submit → จัดซื้อเห็น Quotation ทันที')

add_heading(doc, 'กรณีที่ 2 — ถอนตัว (Decline to Quote)', level=3)
add_step(doc, 1, 'Supplier กดปุ่ม "ถอนตัว / Decline"')
add_step(doc, 2, 'กรอกเหตุผล (บังคับกรอก)')
add_step(doc, 3, 'กด Confirm')
add_step(doc, 4, 'จัดซื้อเห็นสถานะ "ถอนตัว" พร้อมเหตุผล ใน Tab Invite Suppliers')
add_note(doc, 'Supplier สามารถ "ยกเลิกการถอนตัว" ได้ตราบเท่าที่ RFQ ยังอยู่ในสถานะ Published')

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — EVALUATION & AWARD
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. วงจรที่ 4 — ประเมินผลและ Award', level=1)

add_heading(doc, '5A — ดู Quotation ที่ได้รับ', level=2)
add_para(doc, 'แท็บ "Quotations" ใน RFQ Detail แสดง:')
add_bullet(doc, 'Quotation ที่ Supplier Submit แล้ว เรียงตามวันที่')
add_bullet(doc, 'ราคาแต่ละรายการ เปรียบเทียบกับราคาอ้างอิง (Baseline)')
add_bullet(doc, 'Section แยกสำหรับ "Supplier ที่ถอนตัว" พร้อมเหตุผล')

add_heading(doc, '5B — ประเมินและเลือก Winner', level=2)
add_step(doc, 1, 'เปลี่ยน Status เป็น "Evaluation"')
add_step(doc, 2, 'ใช้แท็บ Evaluation เปรียบเทียบราคาจากทุก Supplier เคียงกัน')
add_step(doc, 3, 'คำนวณ Score จาก: ราคา + คุณภาพ + ระยะเวลาส่งมอบ')
add_step(doc, 4, 'Tick "แนะนำให้เลือก (Recommended Winner)" ที่ Quotation ของผู้ชนะ')
add_warning(doc, 'ระบบจะไม่ยอมเปลี่ยน Status เป็น "Awarded" จนกว่าจะมีการ Mark Recommended Winner')

add_heading(doc, '5C — Award', level=2)
add_step(doc, 1, 'เมื่อเลือก Winner แล้ว → เปลี่ยน Status เป็น "Awarded"')
add_step(doc, 2, 'ระบบสร้าง Award Record อัตโนมัติ')
add_step(doc, 3, 'หน้า RFQ Detail แสดง 🏆 ชื่อ Supplier ที่ชนะ + คะแนน')
add_step(doc, 4, 'ดำเนินการต่อด้วยการออก Purchase Order')

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — SUPPORT MODULES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. โมดูลสนับสนุน', level=1)

add_heading(doc, '6A — Supplier Management', level=2)
add_para(doc, 'เมนู Suppliers แสดงรายชื่อ Supplier ทั้งหมดพร้อมตัวกรองหลายมิติ')
add_heading(doc, 'ตัวกรองที่มีในหน้า Supplier List', level=3)
add_table(doc,
    ['ตัวกรอง', 'ตัวเลือก'],
    [
        ['Status', 'draft / submitted / review / approved / rejected / suspended'],
        ['Tier', 'Silver / Gold / Platinum'],
        ['ประเภทใบรับรอง', 'GMP, HACCP, ISO9001, ISO22000, BRCGS, FSSC22000, HALAL, IFS, KOSHER'],
        ['สถานะใบรับรอง', 'ใช้งานได้ (>90 วัน) / ใกล้หมดอายุ (≤90 วัน) / หมดอายุแล้ว'],
    ],
    col_widths=[5, 12]
)

add_heading(doc, 'สีของ Badge ใบรับรองในตาราง', level=3)
add_table(doc,
    ['สี', 'ความหมาย'],
    [
        ['🟢 สีเขียว', 'ใบรับรองยังใช้งานได้ (เกิน 90 วัน)'],
        ['🟡 สีเหลือง', 'ใกล้หมดอายุ (น้อยกว่า 90 วัน)'],
        ['🔴 สีแดง', 'หมดอายุแล้ว — ควรติดตามให้ Supplier ต่ออายุ'],
    ],
    col_widths=[4, 13]
)

add_heading(doc, '6B — Risk Assessment', level=2)
add_para(doc, 'ระบบประเมินความเสี่ยง Supplier ใน 10 หมวดหมู่:')
add_table(doc,
    ['หมวดความเสี่ยง', 'ความหมาย'],
    [
        ['Food Safety', 'ระบบความปลอดภัยด้านอาหาร'],
        ['Quality System', 'ระบบควบคุมคุณภาพ'],
        ['Delivery & Logistics', 'ความน่าเชื่อถือในการส่งมอบ'],
        ['Financial Stability', 'สถานะทางการเงิน'],
        ['Certification Status', 'สถานะใบรับรองมาตรฐาน'],
        ['Food Fraud', 'ความเสี่ยงการปลอมปน'],
        ['Allergen Management', 'การจัดการสารก่อภูมิแพ้'],
        ['Country of Origin Risk', 'ความเสี่ยงตามแหล่งกำเนิดสินค้า'],
        ['Critical Material', 'การพึ่งพาวัตถุดิบวิกฤต'],
        ['NCR History', 'ประวัติ Non-Conformance Report'],
    ],
    col_widths=[6, 11]
)

add_heading(doc, 'การจัดกลุ่ม ABC / XYZ', level=3)
add_table(doc,
    ['มิติ', 'ระดับ A / X', 'ระดับ B / Y', 'ระดับ C / Z'],
    [
        ['ABC (มูลค่า)', 'มูลค่าสูง', 'มูลค่าปานกลาง', 'มูลค่าต่ำ'],
        ['XYZ (ความแปรปรวน)', 'ความต้องการสม่ำเสมอ', 'ความแปรปรวนปานกลาง', 'ความแปรปรวนสูง'],
    ],
    col_widths=[5, 4, 4, 4]
)
add_note(doc, 'Priority Score 1–9 (9 = Critical) คำนวณจาก ABC × XYZ — ใช้ตัดสินใจลำดับความสำคัญในการจัดซื้อ')

add_heading(doc, '6C — Admin Settings', level=2)
add_table(doc,
    ['Setting', 'รายละเอียด'],
    [
        ['Update Cycle (วัน)', 'กำหนดรอบเวลาที่ Supplier ต้องอัพเดทราคา เช่น 90 วัน'],
        ['Hold Period (วัน)', 'ระยะเวลา "ล็อค" หลังส่งราคา เช่น 30 วัน ไม่ต้องส่งใหม่'],
    ],
    col_widths=[5, 12]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. สิทธิ์การใช้งานแต่ละ Role', level=1)
add_table(doc,
    ['ฟีเจอร์', 'Admin', 'Procurement', 'Supplier', 'QA'],
    [
        ['สร้าง/แก้ไข Catalog', '✅', '✅', '❌', '❌'],
        ['กำหนด Target Qty', '✅', '✅', '❌', '❌'],
        ['กำหนด Nominated Supplier', '✅', '❌', '❌', '✅'],
        ['เสนอราคาใน Portal', '❌', '❌', '✅', '❌'],
        ['Export/Import Excel', '✅', '✅', '❌', '❌'],
        ['สร้าง RFQ', '✅', '✅', '❌', '❌'],
        ['เชิญ Supplier ใน RFQ', '✅', '✅', '❌', '❌'],
        ['Submit Quotation', '❌', '❌', '✅', '❌'],
        ['ถอนตัวจาก RFQ', '❌', '❌', '✅', '❌'],
        ['เลือก Recommended Winner', '✅', '✅', '❌', '✅'],
        ['Award RFQ', '✅', '✅', '❌', '❌'],
        ['จัดการ User', '✅', '❌', '❌', '❌'],
        ['Admin Settings', '✅', '❌', '❌', '❌'],
    ],
    col_widths=[7, 2.5, 2.8, 2.5, 2]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Quick Reference — สรุป Flow แบบย่อ', level=1)

add_heading(doc, 'Flow สำหรับจัดซื้อ', level=2)
add_table(doc,
    ['#', 'การดำเนินการ', 'เมนูในระบบ'],
    [
        ['1', 'กำหนด Target Qty ในแต่ละรายการสินค้า', 'Pricelist → Catalog'],
        ['2', 'ส่ง Excel Checklist ให้ Supplier หรือรอกรอกใน Portal', 'Pricelist → ปุ่มส่ง Excel'],
        ['3', 'Import ราคากลับจาก Excel (ถ้าใช้ช่องทาง Excel)', 'Pricelist → Import Quotation'],
        ['4', 'สร้าง RFQ จาก Catalog หรือสร้างใหม่เอง', 'Pricelist → Create RFQ  หรือ  RFQ → New'],
        ['5', 'Publish RFQ → Supplier รับการเชิญ', 'RFQ Detail → Dropdown Status'],
        ['6', 'ติดตามการตอบกลับ / Supplier ถอนตัว', 'RFQ Detail → Invite Suppliers'],
        ['7', 'Close → Evaluation → เลือก Winner', 'RFQ Detail → Evaluation Tab'],
        ['8', 'Award → ดำเนินการออก PO ต่อ', 'RFQ Detail → Dropdown → Awarded'],
    ],
    col_widths=[1.2, 8, 6.5]
)

doc.add_paragraph()
add_heading(doc, 'Flow สำหรับ Supplier', level=2)
add_table(doc,
    ['#', 'การดำเนินการ', 'เมนูในระบบ'],
    [
        ['1', 'Login เข้าระบบด้วย sup-XXXXX@nslfoods.test', 'หน้า Login'],
        ['2', 'เห็น Catalog ที่ตัวเองมีสิทธิ์เสนอราคา', 'Pricelist'],
        ['3', 'กรอกราคาใน Portal (รายการ Open ทุกรายการ)', 'Pricelist → Catalog → กรอกราคา'],
        ['4', 'รับ Invitation จาก RFQ ที่จัดซื้อสร้าง', 'RFQ List'],
        ['5', 'Submit Quotation หรือแจ้ง Decline พร้อมเหตุผล', 'RFQ Detail → Quotations Tab'],
    ],
    col_widths=[1.2, 8, 6.5]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('เอกสารนี้จัดทำโดย SmartProcure System  |  NSL Foods PLC  |  เมษายน 2026')
r.font.name = 'TH Sarabun New'
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.font.italic = True

# ─── Save ──────────────────────────────────────────────────────────────────
out_path = '/Users/golf/Desktop/Projects/smartprocure/SmartProcure_Workflow_Guide.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
